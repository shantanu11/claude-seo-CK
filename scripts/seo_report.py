#!/usr/bin/env python3
"""
Universal SEO Report Generator - converts any analysis output to shareable DOCX/PDF.

Reads markdown reports (from audit, page analysis) or JSON results (from page-batch,
google APIs) and produces professional Word documents or PDFs.

Usage:
    python seo_report.py --input FULL-AUDIT-REPORT.md --format docx
    python seo_report.py --input results.json --format docx
    python seo_report.py --input FULL-AUDIT-REPORT.md --input ACTION-PLAN.md --format docx
    python seo_report.py --input results.json --type batch --format pdf
    python seo_report.py --dir ./output/ --format docx    # auto-find all report files in dir
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Optional


def _detect_input_type(filepath: str) -> str:
    """Detect whether input is markdown or JSON."""
    ext = Path(filepath).suffix.lower()
    if ext == ".json":
        return "json"
    return "markdown"


def _read_inputs(paths: list) -> tuple:
    """Read all input files and return (combined_text, input_type, data_dict)."""
    texts = []
    data = None
    input_type = "markdown"

    for p in paths:
        if not os.path.exists(p):
            print(f"Warning: File not found: {p}", file=sys.stderr)
            continue

        ftype = _detect_input_type(p)
        if ftype == "json":
            input_type = "json"
            with open(p, "r", encoding="utf-8") as f:
                data = json.load(f)
            # Also create a text representation
            texts.append(_json_to_markdown(data, p))
        else:
            with open(p, "r", encoding="utf-8") as f:
                texts.append(f.read())

    return "\n\n---\n\n".join(texts), input_type, data


def _json_to_markdown(data: dict, source: str) -> str:
    """Convert page-batch or google-api JSON results to markdown."""
    lines = []

    # page-batch results
    if "mode" in data and "results" in data:
        mode_label = {
            "page": "Page SEO", "content": "Content/E-E-A-T",
            "geo": "GEO/AI Citation", "all": "Full Analysis"
        }
        lines.append(f"# Batch {mode_label.get(data.get('mode', ''), 'SEO')} Report")
        lines.append(f"\nAnalyzed **{data.get('total', 0)}** URLs on {datetime.now().strftime('%B %d, %Y')}")
        lines.append("")

        s = data.get("summary", {})
        lines.append("## Summary")
        lines.append(f"- **Average Score:** {s.get('avg_score', 0)}/100")
        lines.append(f"- **Analyzed:** {s.get('analyzed', 0)} | **Errors:** {s.get('errors', 0)}")
        lines.append(f"- **Issues:** {s.get('critical_issues', 0)} critical, {s.get('high_issues', 0)} high, "
                      f"{s.get('medium_issues', 0)} medium, {s.get('low_issues', 0)} low")
        lines.append(f"- **Missing title:** {s.get('pages_missing_title', 0)} | "
                      f"**Missing H1:** {s.get('pages_missing_h1', 0)} | "
                      f"**Missing meta desc:** {s.get('pages_missing_meta_desc', 0)}")
        lines.append(f"- **Thin content:** {s.get('pages_thin_content', 0)} | "
                      f"**Noindex:** {s.get('pages_noindex', 0)}")

        if s.get("avg_readability"):
            lines.append(f"- **Avg Readability:** {s['avg_readability']} | "
                          f"**No Author:** {s.get('pages_no_author', 0)} | "
                          f"**No Dates:** {s.get('pages_no_dates', 0)}")
        if s.get("avg_citability"):
            lines.append(f"- **Avg Citability:** {s['avg_citability']}/100 | "
                          f"**Low Citability:** {s.get('pages_low_citability', 0)} | "
                          f"**No Answer-First:** {s.get('pages_no_answer_first', 0)}")
        lines.append("")

        # Results table
        lines.append("## Per-URL Results")
        lines.append("")

        mode = data.get("mode", "page")
        if mode == "all":
            lines.append("| Score | Read | Cite | Words | Issues | URL |")
            lines.append("|------:|-----:|-----:|------:|-------:|-----|")
        elif mode == "content":
            lines.append("| Score | Readability | Words | Issues | URL |")
            lines.append("|------:|------------:|------:|-------:|-----|")
        elif mode == "geo":
            lines.append("| Score | Citability | Words | Issues | URL |")
            lines.append("|------:|-----------:|------:|-------:|-----|")
        else:
            lines.append("| Score | Status | Words | Issues | URL |")
            lines.append("|------:|-------:|------:|-------:|-----|")

        results = sorted(data.get("results", []), key=lambda r: r.get("score") or 0)
        for r in results:
            score = r.get("score", "ERR")
            seo = r.get("seo") or {}
            wc = seo.get("word_count", 0)
            n = len(r.get("issues", []))
            url = r.get("url", "")

            if mode == "all":
                read = (r.get("content") or {}).get("readability_score", 0)
                cite = (r.get("geo") or {}).get("citability_score", 0)
                lines.append(f"| {score} | {read:.0f} | {cite} | {wc} | {n} | {url} |")
            elif mode == "content":
                read = (r.get("content") or {}).get("readability_score", 0)
                lines.append(f"| {score} | {read:.0f} | {wc} | {n} | {url} |")
            elif mode == "geo":
                cite = (r.get("geo") or {}).get("citability_score", 0)
                lines.append(f"| {score} | {cite} | {wc} | {n} | {url} |")
            else:
                status = r.get("status_code", "---")
                lines.append(f"| {score} | {status} | {wc} | {n} | {url} |")
        lines.append("")

        # Detailed per-URL analysis (always shown)
        lines.append("## Detailed Analysis Per URL")
        lines.append("")

        for r in results:
            url = r.get("url", "")
            score = r.get("score", "?")
            seo = r.get("seo") or {}
            content = r.get("content") or {}
            geo = r.get("geo") or {}
            issues = r.get("issues", [])

            lines.append(f"### {url}")
            lines.append(f"**Score: {score}/100**")
            lines.append("")

            # SEO overview
            lines.append("**On-Page SEO:**")
            lines.append(f"- **Title:** {seo.get('title', 'N/A')} ({seo.get('title_length', 0)} chars)")
            lines.append(f"- **Meta Description:** {(seo.get('meta_description') or 'N/A')[:80]}... ({seo.get('meta_description_length', 0)} chars)")
            lines.append(f"- **H1:** {', '.join(seo.get('h1', [])) or 'N/A'}")
            lines.append(f"- **Word Count:** {seo.get('word_count', 0)}")
            lines.append(f"- **Headings:** {seo.get('h1_count', 0)} H1, {seo.get('h2_count', 0)} H2, {seo.get('h3_count', 0)} H3")
            lines.append(f"- **Links:** {seo.get('internal_links', 0)} internal, {seo.get('external_links', 0)} external")
            lines.append(f"- **Images:** {seo.get('images_total', 0)} total, {seo.get('images_missing_alt', 0)} missing alt")
            lines.append(f"- **Schema:** {', '.join(seo.get('schema_types', [])) or 'None'}")
            lines.append(f"- **Canonical:** {seo.get('canonical', 'N/A')}")
            lines.append(f"- **OG Tags:** {'Yes' if seo.get('has_og') else 'No'} | **Twitter Card:** {'Yes' if seo.get('has_twitter_card') else 'No'}")
            lines.append("")

            # Content data (if available)
            if content:
                lines.append("**Content Quality (E-E-A-T):**")
                lines.append(f"- **Readability:** {content.get('readability_score', 0)} ({content.get('reading_level', 'N/A')})")
                lines.append(f"- **Author:** {content.get('author_name') or ('Present' if content.get('has_author') else 'Missing')}")
                lines.append(f"- **Publish Date:** {'Yes' if content.get('has_publish_date') else 'No'} | **Updated Date:** {'Yes' if content.get('has_update_date') else 'No'}")
                lines.append(f"- **External Citations:** {content.get('external_citation_count', 0)}")
                lines.append(f"- **Content-to-HTML Ratio:** {content.get('content_to_html_ratio', 0)}%")
                lines.append(f"- **Paragraphs:** {content.get('paragraph_count', 0)} | **Lists:** {content.get('list_count', 0)} | **Tables:** {content.get('table_count', 0)}")
                lines.append("")

            # GEO data (if available)
            if geo:
                lines.append("**AI Citation Readiness (GEO):**")
                lines.append(f"- **Citability Score:** {geo.get('citability_score', 0)}/100")
                lines.append(f"- **Answer-First:** {'Yes' if geo.get('has_answer_first') else 'No'}")
                lines.append(f"- **Q&A Pairs:** {geo.get('qa_pairs_count', 0)}")
                lines.append(f"- **Entity Clarity:** {geo.get('entity_clarity', 'N/A')}")
                lines.append(f"- **Citable Passages:** {geo.get('citable_passages', 0)} / {geo.get('passage_count', 0)}")
                lines.append(f"- **Stats/Data Points:** {geo.get('stat_claims_count', 0)}")
                lines.append(f"- **Structured Lists:** {'Yes' if geo.get('has_structured_lists') else 'No'} | **Comparison Table:** {'Yes' if geo.get('has_comparison_table') else 'No'}")
                lines.append("")

            # Issues with evidence and verify links
            if issues:
                # Group by severity
                for sev_level in ["critical", "high", "medium", "low"]:
                    sev_issues = [i for i in issues if i.get("severity") == sev_level]
                    if not sev_issues:
                        continue
                    sev_label = {"critical": "CRITICAL", "high": "HIGH", "medium": "MEDIUM", "low": "LOW"}
                    lines.append(f"**{sev_label[sev_level]} Priority Issues:**")
                    lines.append("")
                    for issue in sev_issues:
                        itype = issue.get("type", "page").upper()
                        lines.append(f"**[{sev_label[sev_level]}:{itype}] {issue['issue']}**")
                        if issue.get("found"):
                            lines.append(f"- **Found:** {issue['found']}")
                        if issue.get("expected"):
                            lines.append(f"- **Expected:** {issue['expected']}")
                        if issue.get("fix"):
                            lines.append(f"- **How to fix:** {issue['fix']}")
                        if issue.get("verify"):
                            lines.append(f"- **Verify:** {issue['verify']}")
                        lines.append("")
            else:
                lines.append("**No issues found.**")
                lines.append("")

    # Google API data (psi, gsc, inspection)
    elif any(k in data for k in ("psi", "gsc", "inspection", "crux")):
        lines.append("# Google SEO API Report")
        lines.append(f"\nGenerated on {datetime.now().strftime('%B %d, %Y')}")
        lines.append(f"\nSource: `{source}`")
        lines.append("\n*Use `google_report.py` for full PDF/DOCX with charts.*")

    return "\n".join(lines)


def generate_docx(text: str, output_path: str, title: Optional[str] = None) -> str:
    """Convert markdown text to a professional DOCX document."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        print("Error: python-docx required. Install: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document()
    navy = RGBColor(0x1E, 0x3A, 0x5F)
    dark_gray = RGBColor(0x4A, 0x55, 0x68)
    green = RGBColor(0x2D, 0x6A, 0x4F)
    red = RGBColor(0xC5, 0x30, 0x30)
    muted = RGBColor(0x6B, 0x72, 0x80)

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Styles
    for style_name, size, color in [
        ("Heading 1", 18, navy), ("Heading 2", 14, navy), ("Heading 3", 12, dark_gray)
    ]:
        style = doc.styles[style_name]
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.name = "Calibri"

    style = doc.styles["Normal"]
    style.font.size = Pt(10)
    style.font.name = "Calibri"

    def _add_table_from_md(header_line, separator_line, row_lines):
        """Parse markdown table and add to doc."""
        headers = [c.strip() for c in header_line.strip("|").split("|")]
        cols = len(headers)
        table = doc.add_table(rows=1 + len(row_lines), cols=cols)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shading = cell._element.get_or_add_tcPr()
            elm = shading.makeelement(qn("w:shd"), {qn("w:fill"): "1E3A5F", qn("w:val"): "clear"})
            shading.append(elm)

        # Rows
        for r_idx, line in enumerate(row_lines):
            cells_data = [c.strip() for c in line.strip("|").split("|")]
            for c_idx in range(min(len(cells_data), cols)):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = cells_data[c_idx]
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)

    # Parse markdown line by line
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)

        # Markdown table
        elif stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|--"):
            header_line = stripped
            separator = lines[i + 1].strip()
            row_lines = []
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith("|"):
                row_lines.append(lines[j])
                j += 1
            _add_table_from_md(header_line, separator, row_lines)
            i = j
            continue

        # Bullet list
        elif stripped.startswith("- "):
            text_content = stripped[2:]
            # Bold markers
            p = doc.add_paragraph(style="List Bullet")
            # Parse **bold** within bullets
            parts = re.split(r'(\*\*[^*]+\*\*)', text_content)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(10)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(10)

        # Horizontal rule / separator
        elif stripped == "---":
            doc.add_page_break()

        # Bold paragraph
        elif stripped.startswith("**") and "**" in stripped[2:]:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*[^*]+\*\*)', stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

        # Regular paragraph
        elif stripped:
            doc.add_paragraph(stripped)

        i += 1

    # Footer
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(f"Report generated on {datetime.now().strftime('%B %d, %Y at %H:%M')}")
    run.font.size = Pt(8)
    run.font.color.rgb = muted

    doc.save(output_path)
    return output_path


def generate_pdf(text: str, output_path: str) -> str:
    """Convert markdown text to PDF via HTML intermediate."""
    try:
        from weasyprint import HTML
    except ImportError:
        print("Error: weasyprint required. Install: pip install weasyprint", file=sys.stderr)
        sys.exit(1)

    # Convert markdown to simple HTML
    html_lines = ['<!DOCTYPE html><html><head><meta charset="UTF-8">']
    html_lines.append('<style>')
    html_lines.append('body { font-family: Calibri, sans-serif; font-size: 11pt; margin: 2cm; color: #333; }')
    html_lines.append('h1 { color: #1e3a5f; font-size: 20pt; border-bottom: 2px solid #1e3a5f; padding-bottom: 5px; }')
    html_lines.append('h2 { color: #1e3a5f; font-size: 15pt; margin-top: 20px; }')
    html_lines.append('h3 { color: #4a5568; font-size: 12pt; }')
    html_lines.append('table { border-collapse: collapse; width: 100%; margin: 10px 0; font-size: 9pt; }')
    html_lines.append('th { background: #1e3a5f; color: white; padding: 6px 8px; text-align: left; }')
    html_lines.append('td { border: 1px solid #ddd; padding: 4px 8px; }')
    html_lines.append('tr:nth-child(even) { background: #f9f9f7; }')
    html_lines.append('ul { margin: 5px 0; }')
    html_lines.append('.footer { color: #6b7280; font-size: 8pt; margin-top: 30px; }')
    html_lines.append('</style></head><body>')

    in_table = False
    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("### "):
            html_lines.append(f"<h3>{stripped[4:]}</h3>")
        elif stripped.startswith("## "):
            html_lines.append(f"<h2>{stripped[3:]}</h2>")
        elif stripped.startswith("# "):
            html_lines.append(f"<h1>{stripped[2:]}</h1>")
        elif stripped.startswith("|") and "---" in stripped:
            continue  # Skip separator
        elif stripped.startswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if not in_table:
                html_lines.append("<table><thead><tr>")
                for c in cells:
                    html_lines.append(f"<th>{c}</th>")
                html_lines.append("</tr></thead><tbody>")
                in_table = True
            else:
                html_lines.append("<tr>")
                for c in cells:
                    html_lines.append(f"<td>{c}</td>")
                html_lines.append("</tr>")
        else:
            if in_table:
                html_lines.append("</tbody></table>")
                in_table = False
            if stripped.startswith("- "):
                # Bold in bullets
                content = stripped[2:]
                content = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', content)
                html_lines.append(f"<li>{content}</li>")
            elif stripped == "---":
                html_lines.append('<div style="page-break-after: always;"></div>')
            elif stripped:
                content = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', stripped)
                html_lines.append(f"<p>{content}</p>")

    if in_table:
        html_lines.append("</tbody></table>")

    html_lines.append(f'<p class="footer">Report generated on {datetime.now().strftime("%B %d, %Y at %H:%M")}</p>')
    html_lines.append("</body></html>")

    html_content = "\n".join(html_lines)
    HTML(string=html_content).write_pdf(output_path)
    return output_path


def find_report_files(directory: str) -> list:
    """Auto-discover report files in a directory."""
    patterns = [
        "FULL-AUDIT-REPORT.md", "ACTION-PLAN.md",
        "GOOGLE-API-REPORT-*.md", "results.json",
    ]
    found = []
    d = Path(directory)
    for f in d.iterdir():
        name = f.name
        if name.endswith(".md") and ("REPORT" in name or "PLAN" in name):
            found.append(str(f))
        elif name.endswith(".json") and "result" in name.lower():
            found.append(str(f))
    return sorted(found)


def main():
    parser = argparse.ArgumentParser(
        description="Universal SEO Report Generator - convert any analysis to shareable DOCX/PDF"
    )
    parser.add_argument(
        "--input", "-i", action="append",
        help="Input file(s): markdown report or JSON results. Can specify multiple.",
    )
    parser.add_argument(
        "--dir", "-d",
        help="Auto-find all report files in this directory",
    )
    parser.add_argument(
        "--format", "-f", choices=["docx", "pdf", "both"], default="docx",
        help="Output format (default: docx)",
    )
    parser.add_argument(
        "--output", "-o",
        help="Output file path (auto-generated if not specified)",
    )
    parser.add_argument(
        "--title", "-t",
        help="Custom report title",
    )

    args = parser.parse_args()

    # Collect input files
    inputs = args.input or []
    if args.dir:
        inputs.extend(find_report_files(args.dir))

    if not inputs:
        print("Error: Provide --input file(s) or --dir directory.", file=sys.stderr)
        parser.print_help()
        sys.exit(1)

    text, input_type, data = _read_inputs(inputs)

    if not text.strip():
        print("Error: No content found in input files.", file=sys.stderr)
        sys.exit(1)

    # Generate output filename
    timestamp = datetime.now().strftime("%Y%m%d-%H%M")
    base_name = args.output or f"SEO-Report-{timestamp}"
    if not base_name.endswith((".docx", ".pdf")):
        base_name_clean = base_name
    else:
        base_name_clean = base_name.rsplit(".", 1)[0]

    files = []

    if args.format in ("docx", "both"):
        path = f"{base_name_clean}.docx"
        generate_docx(text, path, title=args.title)
        files.append(path)
        print(f"Generated: {path}")

    if args.format in ("pdf", "both"):
        path = f"{base_name_clean}.pdf"
        generate_pdf(text, path)
        files.append(path)
        print(f"Generated: {path}")

    print(f"\nDone. {len(files)} file(s) generated.")


if __name__ == "__main__":
    main()
